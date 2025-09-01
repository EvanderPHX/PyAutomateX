from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 创建 Word 文档
doc = Document()
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal'].font.size = Pt(12)

# 添加标题
title = doc.add_paragraph("自动化测试开发工程师 面试题")
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].font.size = Pt(16)
title.runs[0].font.bold = True

doc.add_paragraph()  # 空行

# 面试题部分
section_title = doc.add_paragraph("📘 一、面试题（参考答案在题下）")
section_title.runs[0].font.bold = True
section_title.runs[0].font.size = Pt(14)
doc.add_paragraph()

# 面试题列表（题目 + 答案）
interview_questions = [
    {
        "question": "1. 请描述你使用过的自动化测试工具及应用场景。",
        "answer": "参考答案：可描述使用过的工具如 Selenium、Robot Framework、Airtest、QTP 等，并说明在 Web、移动端或桌面应用中的应用场景。"
    },
    {
        "question": "2. 解释回归测试、冒烟测试和兼容性测试的区别。",
        "answer": "参考答案：回归测试：验证新改动未破坏已有功能；冒烟测试：快速验证核心功能是否可用；兼容性测试：验证软件在不同环境或设备上的表现。"
    },
    {
        "question": "3. 请解释面向对象编程中的继承、封装、多态，并举例说明在自动化测试框架设计中如何应用这些概念。",
        "answer": "参考答案：继承：子类可以复用父类方法；封装：隐藏内部实现，提供公共接口；多态：相同接口可对应不同实现。在自动化测试框架中，可通过基类封装通用操作，子类扩展特定功能，实现可复用和可扩展的测试脚本。"
    },
    {
        "question": "4. 请描述函数（Function/Method）的作用，并举例说明你在自动化测试中是如何设计和使用函数的。",
        "answer": "参考答案：函数是实现代码复用和模块化的基本单位。通过函数可以封装特定功能、减少重复代码、提高可维护性。在自动化测试中，可以设计通用操作函数（如登录、点击按钮、数据校验等），在不同测试用例中调用，提高脚本复用性和可读性。"
    },
    {
        "question": "5. 请描述你在测试框架中如何设计和维护自动化测试脚本。",
        "answer": "参考答案：说明脚本结构设计、模块化、可复用性、维护策略、版本管理和代码质量控制等。"
    },
    {
        "question": "6. 如果自动化测试失败，你会如何分析和处理？",
        "answer": "参考答案：分析日志、重现问题、区分环境问题和代码问题、调整脚本、报告缺陷并回归验证。"
    },
    {
        "question": "7. 你如何与开发团队沟通，保证自动化测试与开发流程协同？",
        "answer": "参考答案：定期沟通、使用缺陷管理工具、参与需求评审、CI/CD 流程集成、共享测试报告。"
    }
]

# 将面试题添加到文档
for item in interview_questions:
    # 题目
    p_q = doc.add_paragraph(item["question"])
    p_q.runs[0].font.size = Pt(12)
    p_q.runs[0].font.bold = True
    # 答题空白区域
    doc.add_paragraph("_" * 100)
    # 答案
    p_a = doc.add_paragraph(item["answer"])
    p_a.runs[0].font.size = Pt(12)
    p_a.runs[0].italic = True
    doc.add_paragraph()  # 空行

# 保存 Word 文档
doc.save("自动化测试开发工程师(初级)_面试题.docx")
print("Word 文档已生成: 自动化测试开发工程师_面试题.docx")
