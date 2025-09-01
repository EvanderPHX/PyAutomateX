from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 创建 Word 文档
doc = Document()
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal'].font.size = Pt(12)

# 添加标题
title = doc.add_paragraph("自动化测试开发工程师 笔试答题卡")
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].font.size = Pt(16)
title.runs[0].font.bold = True

doc.add_paragraph()  # 空行

# 笔试题部分
section_title = doc.add_paragraph("📘 一、笔试题（60分）")
section_title.runs[0].font.bold = True
section_title.runs[0].font.size = Pt(14)

doc.add_paragraph()  # 空行

# 选择/简答题
questions = [
    "1. 简述自动化测试的优缺点。（5分）",
    "2. Selenium、Robot Framework、Airtest 的差别是什么？（5分）",
    "3. SQL 查询 employee 表里薪资最高的前三名员工。（5分）",
    "4. 解释回归测试、冒烟测试、兼容性测试。（5分）"
]

for q in questions:
    p = doc.add_paragraph(q)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    # 添加答题空白区域
    doc.add_paragraph(" " * 100)

# 编程题
coding_questions = [
    "5. 编写 Python/Java 函数，输入一个字符串，去掉重复字符，保留首次出现顺序。（10分）如：输入字符：aabbcdeffg，应该输出的字符为：abcdefg",
    "6. 修复下面 Selenium 脚本中的错误：（10分）"
]

for q in coding_questions:
    p = doc.add_paragraph(q)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    doc.add_paragraph(" " * 120)

# Selenium 代码提示
doc.add_paragraph("提示 Selenium 代码：")
selenium_code = [
    "from selenium import webdriver",
    "driver = webdriver.Chrome()",
    "driver.get('http://example.com/login')",
    "driver.find_element_by_id('username').send_keys('testuser')",
    "driver.find_element_by_name('password').sendkeys('123456')",
    "driver.find_element_by_xpath('//button[text()=\"submit\"]').click()",
    "print(driver.title)",
    "driver.quit()"
]
for line in selenium_code:
    p = doc.add_paragraph(line)
    p.style.font.name = 'Consolas'
    p.style.font.size = Pt(10)

# SQL 实操题
sql_questions = [
    "7. 数据表：orders(order_id, customer_id, amount, order_date)；写 SQL 查询每个客户的订单总金额。（10分）",
    "8. 找出订单金额最高的客户 ID。（10分）"
]
for q in sql_questions:
    p = doc.add_paragraph(q)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    doc.add_paragraph(" " * 120)

# 分页插入答案页
doc.add_page_break()
doc.add_paragraph("📘 笔试题参考答案", style='Normal').runs[0].font.bold = True

answers = [
    "1. 自动化测试优缺点：优点：回归快，减少重复工作；覆盖率高；减少人为错误；可集成 CI/CD；缺点：前期投入高；脚本维护成本大；不适合探索性测试。",
    "2. Selenium：Web UI 自动化，多语言支持；Robot Framework：关键字驱动，非程序员可用；Airtest：适合移动端或游戏 UI 自动化。",
    "3. SELECT * FROM employee ORDER BY salary DESC LIMIT 3;",
    "4. 回归测试：验证新代码未破坏已有功能；冒烟测试：快速验证关键功能是否可用；兼容性测试：验证软件在不同环境/设备上的表现。",
    "5. 去重字符串函数（Python）：\n   def remove_duplicates(s):\n       result = ''\n       for c in s:\n           if c not in result:\n               result += c\n       return result\n   print(remove_duplicates('aabbcdeffg'))  # 输出 abcdefg",
    "6. Selenium 脚本修复：\n   driver.find_element('name', 'password').send_keys('123456')\n   driver.find_element('xpath', '//button[@type=\"submit\"]').click()",
    "7. SELECT customer_id, SUM(amount) AS total_amount FROM orders GROUP BY customer_id;",
    "8. SELECT customer_id FROM orders GROUP BY customer_id ORDER BY SUM(amount) DESC LIMIT 1;"
]

for a in answers:
    p = doc.add_paragraph(a)
    p.runs[0].font.size = Pt(12)

# 保存 Word 文档
doc.save("自动化测试开发工程师(初级)_笔试题.docx")
print("Word 文档已生成: 自动化测试开发工程师(初级)_笔试题.docx")
